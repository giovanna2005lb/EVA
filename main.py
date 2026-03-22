#!/usr/bin/env python3
"""
================================================================================
EVA BACKEND — API Web v1.0
================================================================================
Baseado em:
  - revisor_2vara.py v3.0 (lógica de revisão, prompt, pós-processamento)
  - monitorar_pasta.py (estrutura de pastas e fluxo de arquivos)

Integra os dois scripts num backend FastAPI para uso com frontend web.

Endpoints:
  GET    /api/status
  POST   /api/upload/manual       → substitui o manual (opcional)
  DELETE /api/manual              → restaura o manual padrão
  POST   /api/formatar            → envia N minutas, retorna ZIP para download
  GET    /api/download/{job_id}   → baixa o ZIP gerado
================================================================================
"""

import io
import os
import re
import uuid
import json
import zipfile
import time
from pathlib import Path
from datetime import datetime

import anthropic
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ──────────────────────────────────────────────────────────────────────────────
# Configuração
# ──────────────────────────────────────────────────────────────────────────────

BASE_DIR    = Path(__file__).parent
MANUAL_DIR  = BASE_DIR / "manual_usuario"
OUTPUT_DIR  = BASE_DIR / "output"

for d in (MANUAL_DIR, OUTPUT_DIR):
    d.mkdir(exist_ok=True)

MODELO_CLAUDE         = "claude-opus-4-6"
MAX_TOKENS            = 16000
PAUSA_ENTRE_CHAMADAS  = 4          # segundos entre minutas (rate limit)
EXTENSOES_ACEITAS     = {".docx", ".pdf", ".txt"}


# ──────────────────────────────────────────────────────────────────────────────
# Prompt do sistema — Manual de Estilo v3.1 (padrão embutido)
# ──────────────────────────────────────────────────────────────────────────────

PROMPT_PADRAO = """Você é Eva, assistente de revisão linguística da 2ª Vara Cível da Comarca de Parnaíba/PI. Sua função é EXCLUSIVAMENTE revisar gramática, ortografia, pontuação, clareza, estilo e conformidade com o Manual de Estilo v3.1 da Vara — sem qualquer avaliação jurídica ou de mérito.

REGRAS DO MANUAL DE ESTILO v3.1 — CHECKLIST OBRIGATÓRIO:

1. ESTRUTURA
- Estrutura tripartite: 1. Relatório, 2. Fundamentação, 3. Dispositivo
- Transição obrigatória: "É o relatório. Decido." (fecha o relatório, abre a fundamentação)
- Subtítulos numerados na fundamentação: 2.1, 2.2, 2.3...
- Dispositivo numerado: 3.1, 3.2, 3.3...

2. VERBOS DECISÓRIOS NO DISPOSITIVO — REGRA INEGOCIÁVEL
- TODOS os verbos decisórios devem estar na 1ª pessoa do presente do indicativo
- Formas corretas: JULGO, EXTINGO, CONDENO, DECLARO, INDEFIRO, HOMOLOGO, DEFIRO, DETERMINO, REJEITO, FIXO, DECRETO, RECONHEÇO, ACOLHO, ABSOLVO, MANTENHO, REVOGO, CONCEDO, AUTORIZO, CONVERTO, SUSPENDO, CASSO, ANULO, RESSALVO, CONFIRMO, REFORMO, ADMITO
- PROIBIÇÃO ABSOLUTA DE INFINITIVO no dispositivo
- PROIBIÇÃO DE VOZ PASSIVA SINTÉTICA: NUNCA julga-se, extingue-se, condena-se
- Verbos decisórios em CAIXA ALTA
- EXEMPLO PROIBIDO: "3.1. JULGO PROCEDENTE para: 3.2. Declarar... 3.3. Determinar..."
- EXEMPLO CORRETO: "3.1. JULGO PROCEDENTE os pedidos. 3.2. DECLARO a nulidade... 3.3. DETERMINO..."
- Esta regra é A MAIS IMPORTANTE. Seu descumprimento invalida toda a revisão.

3. ATOS DECISÓRIOS vs. PEDIDOS À SECRETARIA
- Atos decisórios: 1ª pessoa (JULGO, CONDENO, DEFIRO) — são do juiz
- Pedidos à Secretaria: impessoal (Intimem-se, Citem-se, Registre-se) — são do cartório
- NUNCA misturar no mesmo item do dispositivo

4. TERMINOLOGIA OBRIGATÓRIA
- "gratuidade da justiça" (com artigo "da", iniciais minúsculas)
- NUNCA "este juízo" — usar "este magistrado" ou construção impessoal
- NUNCA "vieram os autos conclusos" / "compulsando os autos, verifica-se"
- "baixas de estilo" — NUNCA "cautelas de estilo"
- NUNCA "angularização processual"

5. CACOFONIA
- Evitar "como" antes de substantivos que criem sonoridade sexual
- "como autor" → "na qualidade de autor"; "como a autora" → "na pessoa da autora"

6. VOZ PASSIVA SINTÉTICA NA FUNDAMENTAÇÃO
- PROIBIDO: "Verificou-se que", "Reconhece-se", "Constata-se", "Ressalta-se",
  "Destaca-se", "Observa-se", "Impõe-se", "Registre-se que", "Saliente-se que",
  "Note-se que", "Percebe-se que", "Infere-se que", "Depreende-se que",
  "Extrai-se que", "Conclui-se que"
- Converter para forma direta: "Observa-se que a ré..." → "A ré..."
- "Ressalta-se que o prazo..." → "Cumpre ressaltar que o prazo..."

7. COMANDOS DE SECRETARIA — SIMPLICIDADE
- Pedidos à Secretaria devem ser CURTOS e DIRETOS
- "Proceda-se às anotações e registros necessários no sistema" → "Registre-se."
- "Intimem-se as partes para ciência desta deliberação" → "Intimem-se."
- "Expeça-se mandado de citação em face do réu" → "Citem-se."

8. CONCORDÂNCIA COM SUJEITO POSPOSTO
- "Fixou-se os marcos" → "Foram fixados os marcos"
- "Requereu-se os pedidos" → "Foram requeridos os pedidos"

9. FORMATAÇÃO
- Fonte: Times New Roman 14pt
- Recuo: 1,25 cm | Espaçamento: 1,5 | Margens: 3cm sup/esq, 2cm inf/dir
- Texto justificado | Sem espaço extra entre parágrafos

10. CLAREZA E CONCISÃO
- Eliminar redundâncias e frases excessivamente longas
- "Com efeito" confirma (NÃO é "portanto")
- Proibido empilhar: "assim", "desse modo", "portanto", "diante do exposto" em sequência

11. TEMPOS VERBAIS NO RELATÓRIO — REGRA OBRIGATÓRIA
- Verbos que introduzem alegações, argumentos e narrativas das partes SEMPRE no presente do indicativo
- CORRETO: alega, sustenta, destaca, acrescenta, afirma, nega, impugna, requer, argui, narra, relata, pondera, aduz, pugna, colaciona
- ERRADO: alegou, sustentou, destacou, acrescentou, afirmou, negou, impugnou, requereu, arguiu, narrou, relatou, ponderou, aduziu, pugnou, colacionou
- EXCEÇÃO — fatos pontuais e datados permanecem no passado:
  Ex: "ajuizou a ação em 2022", "a decisão de ID 81152694 deferiu a gratuidade", "a requerida apresentou contestação"
- Regra prática: verbos de fala/alegação das partes → PRESENTE; eventos processuais concretos → PASSADO
- Exemplo de erro: "A parte autora alegou que adquiriu..." → CORRETO: "A parte autora alega que adquiriu..."
- Exemplo de erro: "A requerida sustentou que não possui..." → CORRETO: "A requerida sustenta que não possui..."

INSTRUÇÕES DE SAÍDA:
Devolva DUAS seções separadas pelo marcador ===RELATORIO_DE_ERROS===

PRIMEIRA SEÇÃO (antes do marcador): texto da minuta REVISADO e CORRIGIDO.
NÃO use asteriscos (**) para negrito — escreva texto limpo. Verbos decisórios em CAIXA ALTA.
NÃO use formatação markdown de nenhum tipo.

SEGUNDA SEÇÃO (após o marcador): lista de erros no formato:
[TIPO] Descrição do erro → Correção aplicada

Tipos: [L] Linguístico | [F] Formal | [J] Jurídico-formal

Se não houver erros: "Nenhum erro encontrado."

REGRAS ANTI-ALUCINAÇÃO:
- NÃO invente erros.
- NÃO altere mérito jurídico, argumentação ou conteúdo decisório.
- NÃO adicione informações, citações ou jurisprudências.
- NÃO remova conteúdo substantivo — apenas corrija forma e estilo.
- NÃO use markdown na minuta revisada.
"""


def montar_prompt_sistema(manual_customizado: str = "") -> str:
    """Se o usuário enviou um manual, usa-o. Senão usa o padrão."""
    if manual_customizado.strip():
        return (
            "Você é Eva, assistente de revisão linguística da 2ª Vara Cível de Parnaíba/PI.\n\n"
            "## MANUAL DE ESTILO CUSTOMIZADO\n\n"
            + manual_customizado
            + "\n\n---\n\n"
            "INSTRUÇÕES DE SAÍDA:\n"
            "Devolva DUAS seções separadas pelo marcador ===RELATORIO_DE_ERROS===\n\n"
            "PRIMEIRA SEÇÃO: texto da minuta REVISADO. Sem markdown, sem asteriscos.\n"
            "SEGUNDA SEÇÃO: lista de erros no formato:\n"
            "[TIPO] Descrição → Correção aplicada\n"
            "Tipos: [L] Linguístico | [F] Formal | [J] Jurídico-formal\n"
            "Se não houver erros: 'Nenhum erro encontrado.'\n\n"
            "NUNCA invente erros. NUNCA altere mérito jurídico."
        )
    return PROMPT_PADRAO


# ──────────────────────────────────────────────────────────────────────────────
# Extração de texto
# ──────────────────────────────────────────────────────────────────────────────

def extrair_docx(dados: bytes) -> str:
    doc = Document(io.BytesIO(dados))
    return "\n".join(p.text for p in doc.paragraphs)


def extrair_pdf(dados: bytes) -> str:
    try:
        import pdfplumber
        paginas = []
        with pdfplumber.open(io.BytesIO(dados)) as pdf:
            for pag in pdf.pages:
                t = pag.extract_text()
                if t:
                    paginas.append(t)
        return "\n\n".join(paginas)
    except ImportError:
        raise HTTPException(500, "pdfplumber não instalado. Execute: pip install pdfplumber")


def extrair_texto(nome: str, dados: bytes) -> str:
    ext = Path(nome).suffix.lower()
    if ext == ".docx":
        return extrair_docx(dados)
    elif ext == ".pdf":
        return extrair_pdf(dados)
    elif ext == ".txt":
        return dados.decode("utf-8", errors="ignore")
    raise HTTPException(400, f"Formato '{ext}' não suportado. Use .docx, .pdf ou .txt.")


# ──────────────────────────────────────────────────────────────────────────────
# Redes de segurança (pós-processamento) — extraídas do revisor_2vara v3.0
# ──────────────────────────────────────────────────────────────────────────────

INFINITIVO_PARA_1A_PESSOA = {
    "JULGAR": "JULGO", "EXTINGUIR": "EXTINGO", "CONDENAR": "CONDENO",
    "DECLARAR": "DECLARO", "INDEFERIR": "INDEFIRO", "HOMOLOGAR": "HOMOLOGO",
    "DEFERIR": "DEFIRO", "DETERMINAR": "DETERMINO", "REJEITAR": "REJEITO",
    "FIXAR": "FIXO", "DECRETAR": "DECRETO", "RECONHECER": "RECONHEÇO",
    "ACOLHER": "ACOLHO", "ABSOLVER": "ABSOLVO", "MANTER": "MANTENHO",
    "REVOGAR": "REVOGO", "CONCEDER": "CONCEDO", "NEGAR": "NEGO",
    "CONFIRMAR": "CONFIRMO", "REFORMAR": "REFORMO", "CASSAR": "CASSO",
    "ANULAR": "ANULO", "RESSALVAR": "RESSALVO", "ADMITIR": "ADMITO",
    "CONVERTER": "CONVERTO", "SUSPENDER": "SUSPENDO", "AUTORIZAR": "AUTORIZO",
    "ADVERTIR": "ADVIRTO",
}

VERBOS_DECISORIOS = [
    "JULGO", "EXTINGO", "CONDENO", "DECLARO", "INDEFIRO", "HOMOLOGO",
    "DEFIRO", "DETERMINO", "REJEITO", "FIXO", "DECRETO", "RECONHEÇO",
    "ACOLHO", "DESACOLHO", "ABSOLVO", "MANTENHO", "REVOGO", "CONCEDO",
    "NEGO", "CONFIRMO", "REFORMO", "CASSO", "ANULO", "RESSALVO",
    "ADMITO", "INADMITO", "CONVERTO", "SUSPENDO", "AUTORIZO", "ADVIRTO",
    "ARQUIVEM-SE", "REGISTRE-SE", "PUBLIQUE-SE", "INTIMEM-SE", "CITEM-SE",
    "NOTIFIQUEM-SE", "EXPEÇA-SE", "OFICIE-SE",
]

PEDIDOS_SECRETARIA_SET = {
    "intimem-se", "intime-se", "citem-se", "cite-se", "cumpra-se",
    "registre-se", "publique-se", "expeça-se", "oficie-se",
    "notifiquem-se", "notifique-se", "arquivem-se", "arquive-se",
}

# ──────────────────────────────────────────────────────────────────────────────
# Verbos de alegação das partes: passado → presente (usados no Relatório)
# ──────────────────────────────────────────────────────────────────────────────

VERBOS_ALEGACAO_PASSADO_PARA_PRESENTE = [
    (r'\balegou\b', 'alega'),
    (r'\bsustentou\b', 'sustenta'),
    (r'\bdestacou\b', 'destaca'),
    (r'\bacrescentou\b', 'acrescenta'),
    (r'\bafirmou\b', 'afirma'),
    (r'\barguiu\b', 'argui'),
    (r'\bargüiu\b', 'argui'),
    (r'\bnegou\b', 'nega'),
    (r'\bimpugnou\b', 'impugna'),
    (r'\brequereu\b', 'requer'),
    (r'\bpediu\b', 'pede'),
    (r'\bnarrou\b', 'narra'),
    (r'\brelatou\b', 'relata'),
    (r'\bponderou\b', 'pondera'),
    (r'\baduziu\b', 'aduz'),
    (r'\bpugnou\b', 'pugna'),
    (r'\bcolacionou\b', 'colaciona'),
    (r'\bponderoU\b', 'pondera'),
    (r'\basseverou\b', 'assevera'),
    (r'\bdefendeu\b', 'defende'),
    (r'\bpostulou\b', 'postula'),
    (r'\bpleiteou\b', 'pleiteia'),
    (r'\bexplicou\b', 'explica'),
    (r'\besclareceu\b', 'esclarece'),
    (r'\binformou\b', 'informa'),
]

# Verbos de EVENTO PROCESSUAL CONCRETO — NÃO converter (ficam no passado)
EVENTOS_PROCESSUAIS_CONCRETOS = re.compile(
    r'\b(ajuizou|distribuiu|deferiu|indeferiu|juntou|protocolou|'
    r'apresentou|interpôs|recorreu|proferiu|determinou|designou|'
    r'homologou|extinguiu|condenou|absolveu|decretou|citou|intimou|'
    r'penhorou|bloqueou|liberou|remeteu|encaminhou|certificou)\b',
    re.IGNORECASE
)


def limpar_markdown(texto: str) -> str:
    texto = re.sub(r'\*\*(.+?)\*\*', r'\1', texto)
    texto = re.sub(r'\*(.+?)\*', r'\1', texto)
    texto = re.sub(r'^#{1,6}\s*', '', texto, flags=re.MULTILINE)
    texto = re.sub(r'^\s*[-*]\s+', '', texto, flags=re.MULTILINE)
    return texto


def corrigir_infinitivos_dispositivo(texto: str) -> str:
    linhas = texto.split('\n')
    resultado = []
    em_dispositivo = False
    for linha in linhas:
        stripped = linha.strip()
        if re.match(r'^3\.?\s*\.?\s*DISPOSITIVO', stripped, re.IGNORECASE):
            em_dispositivo = True
        if em_dispositivo and re.match(r'^3\.\d+\.?\s', stripped):
            for inf, conj in INFINITIVO_PARA_1A_PESSOA.items():
                linha = re.sub(r'(3\.\d+\.?\s+)' + re.escape(inf) + r'\b', r'\1' + conj, linha)
        elif not em_dispositivo:
            for inf, conj in INFINITIVO_PARA_1A_PESSOA.items():
                linha = re.sub(r'\b' + re.escape(inf) + r'\b', conj, linha)
        resultado.append(linha)
    return '\n'.join(resultado)


def _eh_pedido_secretaria(linha: str) -> bool:
    stripped = linha.strip()
    if not stripped:
        return False
    tokens = [t for t in re.split(r'[.\s]+', stripped) if t]
    return all(t.lower() in PEDIDOS_SECRETARIA_SET for t in tokens)


def _extrair_advertencia(linha: str):
    stripped = linha.strip()
    m = re.match(
        r'[Ff]ica(?:m)?\s+((?:a |o )?(?:parte|autor[a]?|r[eé][ua]?|requerente|requerida?|'
        r'executad[oa]?|embargante|embargad[oa]?))\s+advertid[oa]s?\s+de\s+que\s+(.*)',
        stripped, re.IGNORECASE
    )
    if m:
        return f"ADVIRTO {m.group(1).strip()} de que {m.group(2).strip()}"
    m = re.match(r'[Aa]dvirta-se\s+(.*)', stripped)
    if m:
        return f"ADVIRTO {m.group(1).strip()}"
    return None


def estruturar_dispositivo_simples(texto: str) -> str:
    if re.search(r'^3\.?\s*\.?\s*DISPOSITIVO', texto, re.MULTILINE | re.IGNORECASE):
        return texto
    if re.search(r'^3\.\d+\.?\s', texto, re.MULTILINE):
        return texto

    linhas = texto.split('\n')
    itens, pedidos, preambulo = [], [], []
    encontrou = False

    padrao_verbo = re.compile(
        r'(?:^|(?:Ante\s+o\s+exposto|Diante\s+do\s+exposto|Pelo\s+exposto|'
        r'Isso\s+posto|À\s+vista\s+do\s+exposto)[,.]?\s*)'
        r'(' + '|'.join(re.escape(v) for v in VERBOS_DECISORIOS if not '-' in v) + r')\b',
        re.IGNORECASE
    )

    for linha in linhas:
        stripped = linha.strip()
        if not stripped:
            continue
        if _eh_pedido_secretaria(stripped):
            pedidos.append(stripped)
            continue
        adv = _extrair_advertencia(stripped)
        if adv:
            itens.append(adv)
            encontrou = True
            continue
        if padrao_verbo.search(stripped):
            encontrou = True
            limpo = re.sub(
                r'^(?:Ante\s+o\s+exposto|Diante\s+do\s+exposto|Pelo\s+exposto|'
                r'Isso\s+posto|À\s+vista\s+do\s+exposto)[,.]?\s*',
                '', stripped
            ).strip()
            if limpo:
                itens.append(limpo)
            continue
        preambulo.append(stripped)

    if not encontrou or not itens:
        return texto

    resultado = list(preambulo) + ([''] if preambulo else [])
    resultado += ['3. DISPOSITIVO', '']
    for i, item in enumerate(itens, 1):
        resultado += [f'3.{i}. {item}', '']

    if pedidos:
        cmds = []
        for cmd in pedidos:
            for t in re.split(r'[.\s]+', cmd.strip()):
                if t:
                    t = re.sub(r'^Intime-se$', 'Intimem-se', t, flags=re.IGNORECASE)
                    t = re.sub(r'^Cite-se$', 'Citem-se', t, flags=re.IGNORECASE)
                    if not t.endswith('.'):
                        t += '.'
                    cmds.append(t)
        resultado.append(' '.join(cmds))
    else:
        resultado.append('Intimem-se. Cumpra-se.')

    return '\n'.join(resultado)


def corrigir_voz_passiva_fundamentacao(texto: str) -> str:
    linhas = texto.split('\n')
    resultado = []
    em_fundamentacao = False

    substituicoes = [
        (r'Verificou-se\s+que\s+', ''),
        (r'Verifica-se\s+que\s+', ''),
        (r'Constata-se\s+que\s+', ''),
        (r'Constata-se\s+(?=(?:o |a |os |as ))', ''),
        (r'Reconhece-se\s+que\s+', ''),
        (r'Ressalta-se\s+que\s+', 'Cumpre ressaltar que '),
        (r'Destaca-se\s+que\s+', 'Cumpre destacar que '),
        (r'Destaca-se\s+(?=(?:o |a |os |as ))', 'Cumpre destacar '),
        (r'Observa-se\s+que\s+', ''),
        (r'Impõe-se\s+(?=(?:o |a ))', 'É imperativo '),
        (r'Registre-se\s+que\s+', ''),
        (r'Saliente-se\s+que\s+', 'Cumpre salientar que '),
        (r'Note-se\s+que\s+', ''),
        (r'Percebe-se\s+que\s+', ''),
        (r'Infere-se\s+que\s+', ''),
        (r'Depreende-se\s+que\s+', ''),
        (r'Extrai-se\s+que\s+', ''),
        (r'Conclui-se\s+que\s+', ''),
    ]

    for linha in linhas:
        stripped = linha.strip()
        if re.search(r'É o relatório\.?\s*Decido\.', stripped, re.IGNORECASE):
            em_fundamentacao = True
            resultado.append(linha)
            continue
        if re.match(r'^3\.?\s*\.?\s*DISPOSITIVO', stripped, re.IGNORECASE):
            em_fundamentacao = False
        if re.match(r'^(?:Ante o exposto|Diante do exposto|Pelo exposto|Isso posto)', stripped, re.IGNORECASE):
            em_fundamentacao = False

        if em_fundamentacao and stripped:
            linha_c = linha
            for padrao, sub in substituicoes:
                linha_c = re.sub(padrao, sub, linha_c, flags=re.IGNORECASE)
            linha_c = re.sub(r'^(\s*)([a-záàâãéêíóôõúü])',
                             lambda m: m.group(1) + m.group(2).upper(), linha_c)
            linha_c = re.sub(r'  +', ' ', linha_c)
            resultado.append(linha_c)
        else:
            resultado.append(linha)

    return '\n'.join(resultado)


def corrigir_verbos_relatorio(texto: str) -> str:
    """
    No trecho do Relatório (seção 1), converte verbos de alegação das partes
    do pretérito perfeito para o presente do indicativo.
    Preserva verbos de eventos processuais concretos (ex: ajuizou, deferiu, apresentou).

    Estratégia linha a linha:
    - Só atua dentro da seção 1. RELATÓRIO
    - Para cada linha, verifica se ela contém APENAS verbos de evento concreto
      como predicado principal — se sim, não converte.
    - Caso contrário, aplica as substituições de alegação.
    """
    linhas = texto.split('\n')
    resultado = []
    em_relatorio = False

    for linha in linhas:
        stripped = linha.strip()

        # Detecta início da seção Relatório
        if re.match(r'^1\.?\s*\.?\s*RELATÓRIO', stripped, re.IGNORECASE):
            em_relatorio = True
            resultado.append(linha)
            continue

        # Detecta fim da seção Relatório
        if re.search(r'É o relatório\.?\s*Decido\.', stripped, re.IGNORECASE):
            em_relatorio = False
            resultado.append(linha)
            continue
        if re.match(r'^2\.?\s*\.?\s*FUNDAMENTAÇÃO', stripped, re.IGNORECASE):
            em_relatorio = False
            resultado.append(linha)
            continue

        if em_relatorio and stripped:
            linha_c = linha

            # Só converte se a linha NÃO for dominada por evento processual concreto
            # (heurística: se começa com sujeito processual + verbo concreto, não mexe)
            eh_evento_concreto = bool(re.match(
                r'^(?:A\s+(?:decisão|sentença|decisão\s+de\s+ID|r\.?\s*)?|'
                r'O\s+(?:juízo|magistrado|feito|processo)|'
                r'Os\s+autos|As\s+partes\s+foram)\s+\w+',
                stripped, re.IGNORECASE
            ) and EVENTOS_PROCESSUAIS_CONCRETOS.search(stripped))

            if not eh_evento_concreto:
                for padrao, sub in VERBOS_ALEGACAO_PASSADO_PARA_PRESENTE:
                    linha_c = re.sub(padrao, sub, linha_c, flags=re.IGNORECASE)

            resultado.append(linha_c)
        else:
            resultado.append(linha)

    return '\n'.join(resultado)


def simplificar_comandos_secretaria(texto: str) -> str:
    subs = [
        (r'Proceda-se\s+às?\s+anotaç(?:ão|ões)\s+e\s+registros?\s+necessári[oa]s?\s+'
         r'(?:no sistema|junto ao sistema|no PJe|no sistema PJe)[.\s]*', 'Registre-se.'),
        (r'Proceda-se\s+às?\s+anotaç(?:ão|ões)\s+necessári[oa]s?[.\s]*', 'Registre-se.'),
        (r'Proceda-se\s+aos?\s+registros?\s+necessári[oa]s?[.\s]*', 'Registre-se.'),
        (r'Intimem-se\s+as\s+partes\s+para\s+ciência\s+(?:desta|da\s+presente)\s+'
         r'(?:deliberação|decisão|sentença)[.\s]*', 'Intimem-se.'),
        (r'Expeça-se\s+mandado\s+de\s+citação\s+em\s+face\s+d[oa]s?\s+r[eé][ua]s?[.\s]*', 'Citem-se.'),
        (r'Dê-se\s+ciência\s+(?:às|aos|das|dos)\s+partes?[.\s]*', 'Intimem-se.'),
    ]
    for padrao, sub in subs:
        texto = re.sub(padrao, sub, texto, flags=re.IGNORECASE)
    return texto


def corrigir_concordancia_sujeito_posposto(texto: str) -> str:
    subs = [
        (r'\bFixou-se\s+(os\s)', r'Foram fixados \1'),
        (r'\bFixou-se\s+(as\s)', r'Foram fixadas \1'),
        (r'\bRequereu-se\s+(os\s)', r'Foram requeridos \1'),
        (r'\bRequereu-se\s+(as\s)', r'Foram requeridas \1'),
        (r'\bJuntou-se\s+(os\s)', r'Foram juntados \1'),
        (r'\bJuntou-se\s+(as\s)', r'Foram juntadas \1'),
        (r'\bApresentou-se\s+(os\s)', r'Foram apresentados \1'),
        (r'\bApresentou-se\s+(as\s)', r'Foram apresentadas \1'),
        (r'\bDeterminou-se\s+(os\s)', r'Foram determinados \1'),
        (r'\bDeterminou-se\s+(as\s)', r'Foram determinadas \1'),
    ]
    for padrao, sub in subs:
        texto = re.sub(padrao, sub, texto)
    return texto


def pipeline_pos_processamento(texto: str) -> str:
    """Aplica todas as redes de segurança em sequência."""
    texto = limpar_markdown(texto)
    texto = corrigir_infinitivos_dispositivo(texto)
    texto = estruturar_dispositivo_simples(texto)
    texto = corrigir_verbos_relatorio(texto)               # ← NOVO: tempos verbais no Relatório
    texto = corrigir_voz_passiva_fundamentacao(texto)
    texto = simplificar_comandos_secretaria(texto)
    texto = corrigir_concordancia_sujeito_posposto(texto)
    return texto


# ──────────────────────────────────────────────────────────────────────────────
# Geração de documentos .docx
# ──────────────────────────────────────────────────────────────────────────────

PADRAO_VERBOS_NEGRITO = re.compile(
    r'\b(' + '|'.join(re.escape(v) for v in VERBOS_DECISORIOS) + r')\b'
)


def criar_docx_revisado(texto: str) -> bytes:
    """Gera o .docx da minuta corrigida com formatação da Vara e negrito nos verbos."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.left_margin   = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin  = Cm(2)

    for linha in texto.split('\n'):
        linha = linha.strip()
        if not linha:
            doc.add_paragraph('')
            continue

        para = doc.add_paragraph()
        partes = PADRAO_VERBOS_NEGRITO.split(linha)
        for i, parte in enumerate(partes):
            if not parte:
                continue
            run = para.add_run(parte)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = (i % 2 == 1)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def criar_docx_relatorio(nome_arquivo: str, erros_texto: str) -> tuple[bytes, int]:
    """Gera o .docx do relatório de erros individual. Retorna (bytes, total_erros)."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RELATÓRIO DE REVISÃO")
    r.bold = True
    r.font.size = Pt(14)

    for texto_sub, tamanho in [
        (f"Arquivo: {nome_arquivo}", 11),
        (f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}", 11),
        (f"Modelo: {MODELO_CLAUDE}", 10),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto_sub)
        r.font.size = Pt(tamanho)
        if tamanho == 10:
            r.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Separar erros por categoria
    erros_l, erros_f, erros_j = [], [], []
    for linha in erros_texto.split('\n'):
        linha = linha.strip()
        if not linha:
            continue
        if linha.startswith('[L]'):
            erros_l.append(linha)
        elif linha.startswith('[F]'):
            erros_f.append(linha)
        elif linha.startswith('[J]'):
            erros_j.append(linha)
        elif linha.lower() != 'nenhum erro encontrado.':
            erros_l.append(linha)

    categorias = [
        ("Erros Linguísticos [L]",      erros_l, RGBColor(0, 0, 139)),
        ("Erros Formais [F]",            erros_f, RGBColor(139, 69, 0)),
        ("Erros Jurídico-Formais [J]",   erros_j, RGBColor(139, 0, 0)),
    ]

    total = 0
    for titulo_cat, lista, cor in categorias:
        if lista:
            p = doc.add_paragraph()
            r = p.add_run(titulo_cat)
            r.bold = True
            r.font.color.rgb = cor
            for erro in lista:
                pe = doc.add_paragraph()
                pe.add_run(erro).font.size = Pt(11)
                total += 1

    if total == 0:
        p = doc.add_paragraph()
        r = p.add_run("Nenhum erro encontrado.")
        r.bold = True
        r.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Total de correções: {total}").bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), total


def criar_docx_consolidado(resultados: list) -> bytes:
    """Gera o relatório consolidado do lote."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RELATÓRIO CONSOLIDADO DE REVISÃO")
    r.bold = True
    r.font.size = Pt(16)

    for texto_sub, tamanho in [
        (f"EVA — Projeto Revisor v3.0 · {datetime.now().strftime('%d/%m/%Y %H:%M')}", 11),
        (f"Modelo: {MODELO_CLAUDE}", 10),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto_sub)
        r.font.size = Pt(tamanho)
        if tamanho == 10:
            r.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    total_arq    = len(resultados)
    total_erros  = sum(r.get('total_erros', 0) for r in resultados if r.get('total_erros', 0) >= 0)
    sem_erro     = sum(1 for r in resultados if r.get('total_erros', 0) == 0)
    falhas       = sum(1 for r in resultados if r.get('status') == 'erro')

    for label, val in [
        ("Arquivos processados:",         str(total_arq)),
        ("Total de correções aplicadas:", str(total_erros)),
        ("Arquivos sem erros:",           str(sem_erro)),
        ("Falhas:",                        str(falhas)),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label} {val}").bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("DETALHAMENTO POR ARQUIVO").bold = True
    p.runs[0].font.size = Pt(13)

    for r in resultados:
        doc.add_paragraph()
        p = doc.add_paragraph()
        status_str = "ERRO" if r.get('status') == 'erro' else f"{r.get('total_erros', 0)} correção(ões)"
        p.add_run(f"{r['arquivo']} — {status_str}").bold = True
        if r.get('erros_texto'):
            for linha in r['erros_texto'].split('\n'):
                linha = linha.strip()
                if linha:
                    pe = doc.add_paragraph()
                    pe.add_run(f"  {linha}").font.size = Pt(10)

    # Frequência de tipos de erro
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("ERROS MAIS FREQUENTES").bold = True
    p.runs[0].font.size = Pt(13)

    contagem = {}
    for r in resultados:
        for linha in (r.get('erros_texto') or '').split('\n'):
            linha = linha.strip()
            if linha and linha[:3] in ('[L]', '[F]', '[J]'):
                contagem[linha[:3]] = contagem.get(linha[:3], 0) + 1

    nomes_tipo = {'[L]': 'Linguístico', '[F]': 'Formal', '[J]': 'Jurídico-formal'}
    if contagem:
        for tipo, qtd in sorted(contagem.items(), key=lambda x: -x[1]):
            pe = doc.add_paragraph()
            pe.add_run(f"  {nomes_tipo.get(tipo, tipo)}: {qtd} ocorrência(s)")
    else:
        doc.add_paragraph().add_run("  Nenhum erro encontrado no lote.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# Chamada à API Anthropic
# ──────────────────────────────────────────────────────────────────────────────

def revisar_minuta(texto: str, prompt_sistema: str) -> tuple[str, str]:
    """
    Envia a minuta para a API e retorna (texto_revisado, erros_texto).
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        config_path = BASE_DIR / "config.txt"
        if config_path.exists():
            for linha in config_path.read_text(encoding="utf-8").splitlines():
                linha = linha.strip()
                if linha and not linha.startswith("#"):
                    chave = linha.split("=", 1)[-1].strip()
                    if chave.startswith("sk-ant-"):
                        api_key = chave
                        break

    if not api_key:
        raise HTTPException(500, "ANTHROPIC_API_KEY não configurada. Defina a variável de ambiente ou crie config.txt.")

    cliente = anthropic.Anthropic(api_key=api_key)

    try:
        resposta = cliente.messages.create(
            model=MODELO_CLAUDE,
            max_tokens=MAX_TOKENS,
            system=prompt_sistema,
            messages=[{"role": "user", "content": f"Revise a minuta abaixo:\n\n{texto}"}],
        )
    except anthropic.AuthenticationError:
        raise HTTPException(401, "Chave da API Anthropic inválida.")
    except anthropic.RateLimitError:
        raise HTTPException(429, "Limite da API atingido. Aguarde alguns segundos.")
    except anthropic.APIConnectionError:
        raise HTTPException(502, "Sem conexão com a API Anthropic.")
    except Exception as e:
        raise HTTPException(502, f"Erro na API: {str(e)}")

    texto_resposta = resposta.content[0].text

    if "===RELATORIO_DE_ERROS===" in texto_resposta:
        partes = texto_resposta.split("===RELATORIO_DE_ERROS===", 1)
        return partes[0].strip(), partes[1].strip()
    else:
        return texto_resposta.strip(), "Marcador de erros não encontrado na resposta."


# ──────────────────────────────────────────────────────────────────────────────
# Helpers de manual
# ──────────────────────────────────────────────────────────────────────────────

def obter_manual_usuario() -> str:
    arquivos = sorted(MANUAL_DIR.iterdir(), key=lambda f: f.stat().st_mtime) if MANUAL_DIR.exists() else []
    if arquivos:
        try:
            return extrair_texto(arquivos[-1].name, arquivos[-1].read_bytes())
        except Exception:
            pass
    return ""


def manual_nome() -> str:
    arquivos = sorted(MANUAL_DIR.iterdir(), key=lambda f: f.stat().st_mtime) if MANUAL_DIR.exists() else []
    return arquivos[-1].name if arquivos else "Manual padrão EVA v3.1"


# ──────────────────────────────────────────────────────────────────────────────
# FastAPI
# ──────────────────────────────────────────────────────────────────────────────

app = FastAPI(title="EVA — Revisão de Minutas", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/status")
def status():
    """Verifica o estado do sistema."""
    tem_key = bool(os.environ.get("ANTHROPIC_API_KEY") or (BASE_DIR / "config.txt").exists())
    customizado = any(MANUAL_DIR.iterdir()) if MANUAL_DIR.exists() else False
    return {
        "pronto": tem_key,
        "api_key_configurada": tem_key,
        "manual_customizado": customizado,
        "manual_nome": manual_nome(),
        "modelo": MODELO_CLAUDE,
        "erros": [] if tem_key else ["API key não encontrada. Configure ANTHROPIC_API_KEY ou crie config.txt."],
    }


@app.post("/api/upload/manual")
async def upload_manual(arquivo: UploadFile = File(...)):
    """
    Envia um manual de estilo personalizado (opcional).
    Se não enviado, o sistema usa o Manual de Estilo v3.1 embutido.
    Aceita .docx, .pdf ou .txt.
    """
    ext = Path(arquivo.filename).suffix.lower()
    if ext not in EXTENSOES_ACEITAS:
        raise HTTPException(400, f"Formato '{ext}' não suportado. Use .docx, .pdf ou .txt.")

    dados = await arquivo.read()
    if not dados:
        raise HTTPException(400, "Arquivo vazio.")

    texto = extrair_texto(arquivo.filename, dados)
    if not texto.strip():
        raise HTTPException(400, "Não foi possível extrair texto do arquivo.")

    for f in MANUAL_DIR.iterdir():
        f.unlink()

    (MANUAL_DIR / arquivo.filename).write_bytes(dados)

    return {
        "ok": True,
        "mensagem": f"Manual '{arquivo.filename}' salvo. Será usado nas próximas revisões.",
        "nome": arquivo.filename,
        "caracteres": len(texto),
    }


@app.delete("/api/manual")
def remover_manual():
    """Remove o manual customizado e volta ao padrão (Manual de Estilo v3.1 embutido)."""
    removidos = 0
    for f in MANUAL_DIR.iterdir():
        f.unlink()
        removidos += 1
    return {
        "ok": True,
        "mensagem": "Manual customizado removido. O sistema voltará a usar o Manual de Estilo v3.1 padrão.",
        "arquivos_removidos": removidos,
    }


@app.post("/api/formatar")
async def formatar(arquivos: list[UploadFile] = File(...)):
    """
    Recebe 1 ou mais minutas, processa todas e retorna um ZIP contendo:
      - REVISADO_{nome}.docx   → minuta corrigida com formatação da Vara
      - RELATORIO_{nome}.docx  → relatório de erros individual
      - CONSOLIDADO.docx       → resumo de todo o lote
    """
    if not arquivos:
        raise HTTPException(400, "Nenhum arquivo enviado.")

    for arq in arquivos:
        ext = Path(arq.filename).suffix.lower()
        if ext not in EXTENSOES_ACEITAS:
            raise HTTPException(400, f"'{arq.filename}': formato '{ext}' não suportado.")

    manual_usuario = obter_manual_usuario()
    prompt_sistema = montar_prompt_sistema(manual_usuario)

    job_id = str(uuid.uuid4())
    resultados = []
    arquivos_zip = {}

    for idx, arq in enumerate(arquivos):
        dados = await arq.read()
        nome = arq.filename
        nome_base = Path(nome).stem

        if not dados:
            resultados.append({"arquivo": nome, "status": "erro", "detalhe": "Arquivo vazio.", "total_erros": 0})
            continue

        try:
            texto_original = extrair_texto(nome, dados)

            if len(texto_original.strip()) < 50:
                resultados.append({"arquivo": nome, "status": "erro", "detalhe": "Arquivo muito curto.", "total_erros": 0})
                continue

            texto_revisado, erros_texto = revisar_minuta(texto_original, prompt_sistema)
            texto_revisado = pipeline_pos_processamento(texto_revisado)

            bytes_revisado = criar_docx_revisado(texto_revisado)
            bytes_relatorio, total_erros = criar_docx_relatorio(nome, erros_texto)

            arquivos_zip[f"REVISADO_{nome_base}.docx"]  = bytes_revisado
            arquivos_zip[f"RELATORIO_{nome_base}.docx"] = bytes_relatorio

            resultados.append({
                "arquivo": nome,
                "status": "ok",
                "total_erros": total_erros,
                "erros_texto": erros_texto,
            })

        except HTTPException as e:
            resultados.append({"arquivo": nome, "status": "erro", "detalhe": e.detail, "total_erros": -1, "erros_texto": ""})
        except Exception as e:
            resultados.append({"arquivo": nome, "status": "erro", "detalhe": str(e), "total_erros": -1, "erros_texto": ""})

        if idx < len(arquivos) - 1:
            time.sleep(PAUSA_ENTRE_CHAMADAS)

    bytes_consolidado = criar_docx_consolidado(resultados)
    arquivos_zip[f"CONSOLIDADO_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"] = bytes_consolidado

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for nome_arq, conteudo in arquivos_zip.items():
            zf.writestr(nome_arq, conteudo)
    zip_buf.seek(0)

    zip_path = OUTPUT_DIR / f"{job_id}.zip"
    zip_path.write_bytes(zip_buf.getvalue())

    processados = sum(1 for r in resultados if r["status"] == "ok")
    falhas      = sum(1 for r in resultados if r["status"] == "erro")
    total_erros = sum(r.get("total_erros", 0) for r in resultados if r.get("total_erros", 0) >= 0)

    return {
        "ok": True,
        "job_id": job_id,
        "download_url": f"/api/download/{job_id}",
        "resumo": {
            "total_arquivos": len(arquivos),
            "processados": processados,
            "falhas": falhas,
            "total_correcoes": total_erros,
        },
        "arquivos": resultados,
        "modelo_usado": MODELO_CLAUDE,
        "manual_usado": manual_nome(),
    }


@app.get("/api/download/{job_id}")
def download(job_id: str):
    """Baixa o ZIP com todos os documentos gerados pelo job."""
    zip_path = OUTPUT_DIR / f"{job_id}.zip"
    if not zip_path.exists():
        raise HTTPException(404, "Arquivo não encontrado. O job pode ter expirado — refaça a formatação.")

    nome_download = f"EVA_Revisao_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
    return FileResponse(
        str(zip_path),
        media_type="application/zip",
        filename=nome_download,
    )

from fastapi.staticfiles import StaticFiles
app.mount("/", StaticFiles(directory=str(BASE_DIR / "frontend"), html=True), name="frontend")